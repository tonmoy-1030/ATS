from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse
from django.contrib.auth.decorators import user_passes_test
from candidates.models import Candidate, Offer
from jobs.models import Job, InterviewSchedule, FinalInterviewSchedule
from django.http import HttpResponse
import docx
from django.views.generic import ListView
from django.utils import timezone
from .utils.utils import interview_assessment, create_offer_letter
from io import BytesIO
from datetime import timedelta
from PyPDF2 import PdfMerger
import os
from django.shortcuts import HttpResponse
from employees.models import (Employee, EmployeeConfirmation,
                              TransferOrder, PostingOrder,
                              SalaryInfo)
from .utils.Joining_form import create_joining_form
from .utils.appointment_letter import appointment_letter
from .utils.confirmation_document import confirm_appraisal_paper, confirmation_letter, extension_letter
from .utils.transfer_posting_order import transfer_letter, posting_letter
from .utils.number2text import format_number, convert_to_words
from .forms import (EmployeeFilter, EmployeeConfirmationFilter,
                    EmployeeConfirmationLetterFilter, PostingLetterFilter,
                    TransferLetterFilter, AppointmentLetterFilter, JobOfferFilter, CandidateDetailsForm, JobReportForm)
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .utils.envelope import create_envelope, createNameTag
from django.db import connection
from io import BytesIO
from django.http import HttpResponse
from django.shortcuts import render
from django.db import connection
from django.contrib import messages
from .exporter.exporter import Exporter
from reportlab.lib.pagesizes import landscape, A4, letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import pandas as pd



def generate_reference_number():
    current_year = timezone.now().year
    initial_interview_count = InterviewSchedule.objects.filter(interview_date__year=current_year).count()
    final_interview_count = FinalInterviewSchedule.objects.filter(interview_date__year=current_year).count()
    interview_count = initial_interview_count + final_interview_count
    reference_number = f"Ref: {interview_count:02d}/{current_year}"
    return reference_number

def job_report(request, job_id):
    job = get_object_or_404(Job, pk=job_id)
    total_candidates = Candidate.objects.filter(job=job).count()
    total_present = Candidate.objects.filter(job=job, attendance_status='Present').count()
    total_forwarded = Candidate.objects.filter(job=job, initial_interview_status='Forwarding for the next Interivew').count()
    total_requirements = job.no_of_position
    interview_schedule = InterviewSchedule.objects.filter(job=job).first()
    if interview_schedule:
        interviewers = interview_schedule.interviewer.all()
    else:
        interviewers = []
        
    
    # Calculate percentages
    percentage_present = (total_present / total_candidates) * 100 if total_candidates > 0 else 0
    percentage_forwarded = (total_forwarded / total_candidates) * 100 if total_candidates > 0 else 0
    
    # Get the reference number
    reference_number = generate_reference_number()
    
    context = {
        'job': job,
        'total_candidates': total_candidates,
        'total_present': total_present,
        'total_forwarded': total_forwarded,
        'total_requirements': total_requirements,
        'percentage_present': round(percentage_present, 2),
        'percentage_forwarded': round(percentage_forwarded, 2),
        'job_id': job.id,
        'reference_number': reference_number,
        'interviewers':interviewers
    }
    return render(request, 'documents/job_report.html', context)


def job_report_document(request, job_id):
    # Load the Word document template
    doc = docx.Document(r'..\ats\templates_for_documents\Job-Report.docx')
    job = get_object_or_404(Job, pk=job_id)
    total_candidates = Candidate.objects.filter(job=job).count()
    total_present = Candidate.objects.filter(job=job, attendance_status='Present').count()
    total_forwarded = Candidate.objects.filter(job=job, initial_interview_status='Forwarding for the next Interivew').count()
    total_requirements = job.no_of_position
    interview_schedule = InterviewSchedule.objects.filter(job=job).first()
    if interview_schedule:
        interviewers = interview_schedule.interviewer.all()
    else:
        interviewers = []
        
    
    # Calculate percentages
    percentage_present = (total_present / total_candidates) * 100 if total_candidates > 0 else 0
    percentage_forwarded = (total_forwarded / total_candidates) * 100 if total_candidates > 0 else 0

    reference_number = generate_reference_number()

    # Define placeholders and corresponding values
    placeholders = {
        '{REF}': reference_number,
        '{JOB_TITLE}': job.job_title,
        '{TOTAL_CANDIDATES}': str(total_candidates),
        '{TOTAL_PRESENT}': f"{total_present} ({percentage_present}%)",
        '{TOTAL_FORWARDED}': f"{total_forwarded} ({percentage_forwarded}%)",
        '{TOTAL_REQUIREMENTS}': str(total_requirements),
        '{INTERVIEWER}':'\n'.join(f"{i+1}. {interviewer.name}" for i, interviewer in enumerate(interviewers)),
    }

    for p, v in placeholders.items():
        for paragraph in doc.paragraphs:
            if p in paragraph.text:
                paragraph.text = paragraph.text.replace(p, v)

    # Replace placeholders with actual values in table cells
    for row in doc.tables[0].rows:
        for cell in row.cells:
            for p, v in placeholders.items():
                if p in cell.text:
                    cell.text = cell.text.replace(p, v)

    # Save the modified document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=job_report.docx'
    doc.save(response)
    return response

def Attendance_sheet(request, interview_id):
    initial_interview = get_object_or_404(InterviewSchedule, pk=interview_id)
    candidate_list = Candidate.objects.filter(interview_schedule=initial_interview, invitation_status='Confirmed')
    context = {
        'interview':initial_interview,
        'candidate_list':candidate_list
    }
    return render (request, 'documents/interview_attendance.html', context=context)

def generate_candidate_assessment(request, interview_id):
    interview_schedule = get_object_or_404(InterviewSchedule, pk=interview_id)
    candidate_list = Candidate.objects.filter(interview_schedule=interview_schedule, attendance_status="Present").distinct()
    merged_buffer = BytesIO()
    merger = PdfMerger()
    for candidate in candidate_list:    
        buffer = BytesIO()
        if not hasattr(candidate, 'details'):
            education_qualification = ""
            age = ""
            total_exp = ""
        else:
            education_qualification = candidate.details.degree_name
            age = str(timezone.now().year - candidate.details.date_of_birth.year) + " years"
            total_exp = candidate.details.total_experience
        candidate_data = {
        'date': candidate.interview_schedule.interview_date,
        'name': candidate.name,
        'applied_for': candidate.interview_schedule.job.first().job_title,
        'educational_qualification': education_qualification,
        'age':age,
        'unit': candidate.interview_schedule.job.first().unit.short_name,
        'total_experience': total_exp
    }
        interview_assessment(candidate_data=candidate_data, pdf_path=buffer)
        buffer.seek(0)
        merger.append(buffer)
    merger.write(merged_buffer)
    merged_buffer.seek(0)
    response = HttpResponse(merged_buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="candidate_assessments.pdf"'
    
    return response

class OfferedCandidate(ListView):
    template_name = 'documents/offered_candidates.html'
    context_object_name = 'offered_candidate'
    
    def get_queryset(self):
        return Candidate.objects.filter(offer__isnull=False, offer__joining_date__gte=timezone.now().date()) 

def offer_letter_generator(request):
    if request.method == "POST":
        offered_candidate_ids = request.POST.getlist('offered_candidates')
        offers = Offer.objects.filter(id__in=offered_candidate_ids)
        merged_buffer = BytesIO()
        merger = PdfMerger()
        for offer in offers:
            buffer = BytesIO()
        
            reference = f"T.K./{offer.job.unit.short_name}/HR/OFR/{offer.reference_number:02d}/{offer.offer_date.month:02d}/{timezone.now().year}"            
            candidate_data={
            'ref': reference,
            'date': timezone.now().date,
            'name': offer.candidate.name,
            'designation': offer.offered_designation,
            'joining_date': offer.joining_date,
            'unit': offer.job.unit.factory,
            'vill': offer.candidate.details.permanent_vill,
            'po': offer.candidate.details.permanent_po,
            'ps': offer.candidate.details.permanent_ps,
            'dist': offer.candidate.details.permanent_dist,
            'location':offer.job.unit.floor_location,
        }
    
            create_offer_letter(candidate_data=candidate_data, pdf_file=buffer)
            buffer.seek(0)
            merger.append(buffer)
        
        merger.write(merged_buffer)
        merged_buffer.seek(0)
       
        response = HttpResponse(merged_buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'inline; filename="offer_letter.pdf"'  
        return response
    
    return HttpResponse("connection ok")



def employee_list(request):
    employee_filter = EmployeeFilter(request.GET, queryset=Employee.objects.filter(details__isnull=False).order_by('-DOJ'))
    
    paginator = Paginator(employee_filter.qs, 10)  
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    return render(request, 'documents/employee_joining_form.html', {'filter': employee_filter, 'page_obj': page_obj})

def JobOfferList(request):
    employee_filter = JobOfferFilter(request.GET, queryset=Offer.objects.filter(joining_date__gte=timezone.now().date(), offer_status="accepted"))
    
    paginator = Paginator(employee_filter.qs, 10)  
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    return render(request, 'documents/offered_candidates.html', {'filter': employee_filter, 'page_obj': page_obj})

def employee_confirmation_list(request):
    employee_filter = EmployeeConfirmationFilter(request.GET, queryset=Employee.objects.all().order_by('id')) 
    
    paginator = Paginator(employee_filter.qs, 10)  
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    return render(request, 'documents/employee_confirmation_list.html', {'filter': employee_filter, 'page_obj': page_obj})


# Confirmation Employee List
def employee_confirmation_letter_list(request):
    employee_filter = EmployeeConfirmationLetterFilter(request.GET, queryset=EmployeeConfirmation.objects.all().\
        order_by('employee__unit')) 
    paginator = Paginator(employee_filter.qs, 10)  
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    return render(request, 'documents/employee_confirmation_letter_list.html', {'filter': employee_filter, 'page_obj': page_obj})



def joining_form_generator(request):
    settings_dir = os.path.dirname(__file__)
    project_root = os.path.abspath(os.path.dirname(settings_dir))
    benevolent_fund = os.path.join(project_root, 'documents/utils/MM. Benevolent Fund Application Form.pdf')
    
    merged_buffer = BytesIO()
    merger = PdfMerger()
    
    if request.method == "POST":
        employee_ids = request.POST.getlist('employees_list')
        employee_list = Employee.objects.filter(id__in=employee_ids)
        
        for employee in employee_list:
            if not hasattr(employee, 'details'):
                return HttpResponse(f"{employee.name} has no details information available, please update information and try again!")
            else:
               
                employee_data = {
                    "eid": employee.EID,
                    "name": employee.name,
                    "doj": employee.DOJ,
                    "designation": employee.designation,
                    "dept": employee.department,
                    "company": employee.unit.factory,
                    'unit': employee.unit.name,
                    'maritial_status':  employee.details.marital_status,
                    "father_name": employee.details.father_name,
                    "mother_name": employee.details.mother_name,
                    "spouse_name": employee.details.spouse_name,
                    "present_address": f"Vill: {employee.details.present_vill}, P.O: {employee.details.present_po}, \
                    P.S: {employee.details.present_ps}, Dist: {employee.details.present_dist}",
                    "permanent_address": f"Vill: {employee.details.permanent_vill}, P.O: {employee.details.permanent_po}, \
                    P.S: {employee.details.permanent_ps}, Dist: {employee.details.permanent_dist}",
                    "personal_mobile_no": employee.mobile_no,
                    "religion": employee.details.religion,
                    "blood_group": employee.details.blood_group,
                    "NID": employee.details.nid,
                    "tin": employee.details.tin,
                    "dob": employee.details.date_of_birth,
                    "personal_email": employee.email,
                    "son": employee.details.no_of_son,
                    "daughter": employee.details.no_of_daughter,
                    "emergency_contact_person": employee.details.emergency_contact_person,
                    "emergency_contact_no": employee.details.emergency_contact_no,
                    "emergency_contact_address": employee.details.emergency_person_address,
                    "emer_relation_with_employee":  employee.details.emer_relation_with_employee,
                    "highest_education": employee.details.highest_degree,
                    "high_institution": employee.details.institution_highest_degree,
                    "passing_year": employee.details.passing_year_highest_degree,
                    "professional_degree": employee.details.professional_degree,
                    "pro_institution": employee.details.institution_professional_degree,
                    "pro_passing_year": employee.details.passing_year_professional_degree,
                    "nominee_name": employee.details.nominee_name,
                    "nominee_father": employee.details.nominee_father_name,
                    "nominee_mother": employee.details.nominee_mohter_name,
                    "relation_with_employee": employee.details.relation_with_employee,
                    'nominee_mobile': employee.details.nominee_mobile_no,
                    "nominee_nid": employee.details.nominee_nid,
                    "nominee_address":  f"Vill: {employee.details.nominee_vill}, P.O: {employee.details.nominee_po}, \
                    P.S: {employee.details.nominee_ps}, Dist: {employee.details.nominee_dist}",
                }

                buffer = BytesIO()
                create_joining_form(employee_data=employee_data, pdf_file=buffer)
                buffer.seek(0)
                merger.append(buffer)
                if employee.unit.name != "Prime Pusti Limited" and employee.unit.name != "Prime Cosmetics Limited":
                    with open(benevolent_fund, 'rb') as bf_file:
                        merger.append(bf_file)

        merger.write(merged_buffer)
        merged_buffer.seek(0)
    
        response = HttpResponse(merged_buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'inline; filename="joining-forms.pdf"'
        return response

    return HttpResponse("Please select employee to generate joining forms.")


def confirm_appraisal_form(request):
    merged_buffer = BytesIO()
    merger = PdfMerger()

    if request.method == "POST":
        employee_ids = request.POST.getlist('employees_list')
        employee_list = Employee.objects.filter(id__in=employee_ids) 
        for employee in employee_list:
            employee_data = {
                "EID": employee.EID,
                "name":employee.name,
                "designation": employee.designation,
                "department": employee.department,
                "unit": employee.unit.short_name,
                "location": employee.job_location,
                "joining_date": employee.DOJ,
                "confirmation_date": employee.confirmation_date,
                
            }
            buffer = BytesIO()
            confirm_appraisal_paper(employee_date=employee_data, pdf_file=buffer)
            buffer.seek(0)
            merger.append(buffer)
    merger.write(merged_buffer)
    merged_buffer.seek(0)
    
    response = HttpResponse(merged_buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="confirmation_appraisal_paper.pdf"'
    return response



def generate_confirmation_letter(request):
    merged_buffer = BytesIO()
    merger = PdfMerger()
    
    if request.method == "POST":
        employee_ids = request.POST.getlist('employees_list')
        employee_list = EmployeeConfirmation.objects.filter(id__in=employee_ids)
        
        for employee in employee_list:
                        
            # Determine status reference
            st = "CON" if employee.status == "Confirmation" else "EXT"
            
            # Generate reference number
            reference = f"T.K./{employee.employee.unit.short_name}/HR/{st}/{employee.reference_number:02d}/{employee.issue_date.month:02d}/{timezone.now().year}"
            
            # Collect employee data
            employee_data = {
                "ref": reference,
                "issue_date": employee.issue_date,
                "EID": employee.employee.EID,
                "name": employee.employee.name,
                "designation": employee.employee.designation,
                "new_designation": employee.new_designation,
                "department": employee.employee.department,
                "unit": employee.employee.unit,
                "location": employee.employee.job_location,
                "development_area": employee.development_area,
                "effective_date": employee.effective_date,
                "business_director": employee.employee.unit.business_director
            }
            
            # Generate the appropriate letter
            buffer = BytesIO()
            if employee.status == "Confirmation":
                confirmation_letter(employee_data=employee_data, pdf_file=buffer)
            else:
                extension_letter(employee_data=employee_data, pdf_file=buffer)
            
            buffer.seek(0)
            merger.append(buffer)
    
    merger.write(merged_buffer)
    merged_buffer.seek(0)
    
    response = HttpResponse(merged_buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="confirmation_letter.pdf"'
    return response



def transfer_letter_list(request):
    employee_filter = TransferLetterFilter(request.GET, queryset=TransferOrder.objects.all().order_by('issue_date')) 
    paginator = Paginator(employee_filter.qs, 10)  
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    return render(request, 'documents/transfer_order_list.html', {'filter': employee_filter, 'page_obj': page_obj})

def posting_letter_list(request):
    employee_filter = PostingLetterFilter(request.GET, queryset=PostingOrder.objects.all().order_by('issue_date')) 
    paginator = Paginator(employee_filter.qs, 10)  
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    return render(request, 'documents/posting_order_list.html', {'filter': employee_filter, 'page_obj': page_obj})


def generate_transfer_letter(request):
    merged_buffer = BytesIO()
    merger = PdfMerger()
    buffer = BytesIO()

    if request.method == "POST":
        employee_ids = request.POST.getlist('employees_list')
        transfer_list = TransferOrder.objects.filter(id__in=employee_ids)
        
        for transfer in transfer_list:
            
            reference = f"T.K./{transfer.employee.unit.short_name}/HR/TFR/{transfer.reference_number:02d}/{transfer.issue_date.month:02d}/{timezone.now().year}"            
           
            transfer_data = {
                "ref": reference,
                "issue_date": transfer.issue_date,
                "EID": transfer.employee.EID,
                "name":transfer.employee.name,
                "designation": transfer.employee.designation,
                "new_designation": transfer.new_designation,
                "unit": transfer.employee.unit.name,
                "current_location": transfer.current_job_location,
                "current_type": transfer.current_location_type,
                "current_region/zone": transfer.current_under_region_zone,
                "current_region/zone_type": transfer.current_under_region_zone_type,
                "new_location": transfer.new_job_location,
                "new_type": transfer.new_location_type,
                "new_region/zone": transfer.new_under_region_zone,
                "new_region/zone_type": transfer.new_under_region_zone_type,
                "new_designation": transfer.new_designation,
                "report_to": transfer.report_to,
                "effective_date": transfer.effective_date,
                "signature": transfer.employee.unit.responsible_hr_manager_name,
                "signature_designation":transfer.employee.unit.responsible_hr_manager_designation,
                "business_director":transfer.employee.unit.business_director
                
            }
            
            buffer = BytesIO()
            transfer_letter(transfer_data=transfer_data, pdf_file=buffer)
            buffer.seek(0)
            merger.append(buffer)
            
    merger.write(merged_buffer)
    merged_buffer.seek(0)
    response = HttpResponse(merged_buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="transfer_letter.pdf"'
    return response


def generate_posting_letter(request):
    merged_buffer = BytesIO()
    merger = PdfMerger()
    buffer = BytesIO()

    if request.method == "POST":
        employee_ids = request.POST.getlist('employees_list')
        posting_list = PostingOrder.objects.filter(id__in=employee_ids)
        
        for posting in posting_list:
        
            reference = f"T.K./{posting.employee.unit.short_name}/HR/PO/{posting.reference_number:02d}/{posting.issue_date.month:02d}/{timezone.now().year}"            
           
            posting_data = {
                "ref": reference,
                "issue_date": posting.issue_date,
                "EID": posting.employee.EID,
                "name":posting.employee.name,
                "designation": posting.employee.designation,
                "unit": posting.employee.unit.name,
                "new_location": posting.new_job_location,
                "new_type": posting.new_location_type,
                "new_region/zone": posting.new_under_region_zone,
                "new_region/zone_type": posting.new_under_region_zone_type,
                "report_to": posting.report_to,
                "effective_date": posting.effective_date,
                "signature": posting.employee.unit.responsible_hr_manager_name,
                "signature_designation": posting.employee.unit.responsible_hr_manager_designation,
                "business_director":posting.employee.unit.business_director
                
            }
            buffer = BytesIO()
            posting_letter(posting_data, pdf_file=buffer)
            buffer.seek(0)
            merger.append(buffer)
            
    merger.write(merged_buffer)
    merged_buffer.seek(0)
    response = HttpResponse(merged_buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="Posting_Letter.pdf"'
    return response

def is_admin(user):
    return user.is_authenticated and user.is_staff

@user_passes_test(is_admin, login_url='/accounts/login', redirect_field_name='next')
# Employee List for Appointment Letter
def appointment_letter_list(request):
    
    employee_filter = AppointmentLetterFilter(request.GET, queryset=SalaryInfo.objects.all().order_by('-issue_date'))
    paginator = Paginator(employee_filter.qs, 10)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    context = {
        'filter': employee_filter,
        'page_obj': page_obj,
    }    
    return render(request, 'documents/employee_appointment_letter.html', context=context)



def generate_appointment_letter(request):
    merged_buffer = BytesIO()
    merger = PdfMerger()
    buffer = BytesIO()

    if request.method == "POST":
        employee_ids = request.POST.getlist('employees_list')
        appointment_list = SalaryInfo.objects.filter(id__in=employee_ids)
        
        for appointment in appointment_list:
            
            # Determine Policy
            if appointment.employee.unit.id == 3:
                policy = "Samuda Group"
            elif appointment.employee.unit.id == 4:
                policy = "Samuda Group"   
            else:
                policy = "T.K. Group"
                
            # For Appointment Letter Template    
            if appointment.employee.unit.name == appointment.employee.unit.factory:
                unit = f"in {appointment.employee.unit.name}"
            else: 
                unit = f"for {appointment.employee.unit.name} in {appointment.employee.unit.factory}"
                 
            reference = f"T.K./{appointment.employee.unit.short_name}/HR/APPT/{appointment.reference_number:02d}/{appointment.issue_date.month:02d}/{timezone.now().year}"            
            
            appointment_data = {
                "ref": reference,
                "issue_date": appointment.issue_date,
                "EID": appointment.employee.EID,
                "name":appointment.employee.name,
                "designation": appointment.employee.designation,
                "unit": unit,
                "permanent_vill":appointment.employee.details.permanent_vill,
                "permanent_PO":appointment.employee.details.permanent_po,
                "permanent_PS":appointment.employee.details.permanent_ps,
                "permanent_dist":appointment.employee.details.permanent_dist,
                'floor_location':appointment.employee.unit.floor_location,
                "policy":policy,
                'location':appointment.place_of_posting,
                "DOJ": appointment.employee.DOJ,
                "salary": format_number(appointment.salary),
                "in_word": convert_to_words(appointment.salary),
                "report_to": appointment.report_to,
                "director_signature": appointment.employee.unit.business_director,
                "CC1": appointment.employee.unit.business_director,
                "CC2": appointment.CC2,
                "CC3": appointment.CC3,
                "CC4": appointment.CC4,
                "CC5": appointment.CC5,
                "CC6": appointment.CC6,
                "CC7": appointment.CC7,
            }
            
            buffer = BytesIO()
            appointment_letter(appointment_data, pdf_file=buffer)
            buffer.seek(0)
            merger.append(buffer)
            
    merger.write(merged_buffer)
    merged_buffer.seek(0)
    response = HttpResponse(merged_buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="appointment_Letter.pdf"'
    return response

def employee_list_envelope(request):
    employee_filter = EmployeeFilter(request.GET, queryset=Employee.objects.all().order_by('id')) 
    
    paginator = Paginator(employee_filter.qs, 10)  
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    return render(request, 'documents/employee_envelope_form.html', {'filter': employee_filter, 'page_obj': page_obj})


def envelopePrinting(request):
    merged_buffer = BytesIO()
    merger = PdfMerger()
    buffer = BytesIO()


    if request.method == "POST":
        employee_ids = request.POST.getlist('employees_list')
        employee_list = Employee.objects.filter(id__in=employee_ids)
        
        for employee in employee_list:
            employee_data = {
                
                "EID": employee.EID,
                "name":employee.name,
                "designation": employee.designation,
            }
            
            buffer = BytesIO()
            create_envelope(file_path=buffer, employee_info=employee_data)
            buffer.seek(0)
            merger.append(buffer)
            
    merger.write(merged_buffer)
    merged_buffer.seek(0)
    response = HttpResponse(merged_buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="envelope_.pdf"'
    return response

def employee_list_name_tag(request):
    employee_filter = EmployeeFilter(request.GET, queryset=Employee.objects.all().order_by('id')) 
    
    paginator = Paginator(employee_filter.qs, 10)  
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    return render(request, 'documents/employee_name_tag_form.html', {'filter': employee_filter, 'page_obj': page_obj})


def nameTagPrinting(request):
    merged_buffer = BytesIO()
    merger = PdfMerger()
    buffer = BytesIO()


    if request.method == "POST":
        employee_ids = request.POST.getlist('employees_list')
        employee_list = Employee.objects.filter(id__in=employee_ids)
        
        for employee in employee_list:
            employee_data = {
                
                "EID": employee.EID,
                "name":employee.name,
                "designation": employee.designation,
                'doj':employee.DOJ
            }
            
            buffer = BytesIO()
            createNameTag(file_path=buffer, employee_info=employee_data)
            buffer.seek(0)
            merger.append(buffer)
            
    merger.write(merged_buffer)
    merged_buffer.seek(0)
    response = HttpResponse(merged_buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="envelope_.pdf"'
    return response

def candidateReport(request):
    if request.method == "POST":
        form = CandidateDetailsForm(request.POST)
        if form.is_valid():
            from_date = form.cleaned_data.get("from_date")
            to_date = form.cleaned_data.get("to_date")
            unit = form.cleaned_data.get('unit')
            if len(unit) > 1:
                unit = tuple(unit)
            else:
                for i in unit:
                    unit = i
                unit = f'("{unit}")'
            
            
            # Ensure from_date and to_date are not None
            if not from_date or not to_date:
                return render(request, "documents/candidateDetails.html", {
                    'form': form,
                    'error': 'Please provide both from and to dates.'
                })
            
            # Format the dates for SQL
            from_date_str = from_date.strftime('%Y-%m-%d')
            to_date_str = (to_date + timedelta(days=1)).strftime('%Y-%m-%d')
            
            query = f"""
                    SELECT 
                        ROW_NUMBER() OVER (ORDER BY jobs_businessunit.name DESC, jobs_job.job_title DESC) AS SL,
                        cc.name AS Name,
                        cc.mobile AS Mobile,
                        DATE(js1.interview_date) AS Interview_Date,
                        jobs_job.job_title AS Job_Title,
                        jobs_businessunit.short_name AS Unit,
                        Null as 'Arrival Time',
                        Null as Signature
                    FROM
                        jobs_interviewschedule js1
                            JOIN
                        candidates_candidate cc ON js1.id = cc.interview_schedule_id
                            JOIN
                        jobs_interviewschedule_job jsj ON js1.id = jsj.interviewschedule_id
                            JOIN
                        jobs_job ON jsj.job_id = jobs_job.id
                            JOIN
                        jobs_businessunit on jobs_job.unit_id = jobs_businessunit.id

                    WHERE
                        js1.interview_date BETWEEN '{from_date_str}' AND '{to_date_str}' and jobs_businessunit.id in {unit}
                    order by 
                    jobs_businessunit.name desc,
                    jobs_job.job_title desc;       
            """

            with connection.cursor() as cursor:
                cursor.execute(query)
                results = cursor.fetchall()
                description = cursor.description
                if results !=():
                    headers = [d[0] for d in description]
                    
                    output = Exporter().ExcelExporter(title='Candidate list' , headers=headers, body=results)
                    output.seek(0)

                    response = HttpResponse(output, content_type="application/vnd.ms-excel")
                    response["Content-Disposition"] = 'attachment; filename="Candidate_Report.xlsx"'
                    return response
                else:
                    messages.warning(request, "Empty Results")
                    form = CandidateDetailsForm()                 
                    
    else:
        form = CandidateDetailsForm()

    return render(request, template_name="documents/reportCandidateDetails.html", context={'form': form})
 
        
@user_passes_test(is_admin, login_url='/login', redirect_field_name='next')

def employeeList(request):
    
    query = '''
        SELECT DISTINCT employees_employee.eid,
                employees_employee.name,
                employees_employee.designation,
                employees_employee.department,
                employees_employee.DOJ,
                candidates_candidatesdetails.date_of_birth,
                employees_employee.mobile_no,
                employees_employee.email,
                candidates_candidatesdetails.blood_group,
                employees_salaryinfo.salary,
                employees_salaryinfo.report_to,
                jobs_businessunit.name,
                candidates_candidatesdetails.nid,
                employees_employeedetails.tin,
                employees_employeedetails.father_name,
                employees_employeedetails.mother_name,
                candidates_candidatesdetails.permanent_vill,
                candidates_candidatesdetails.permanent_po,
                candidates_candidatesdetails.permanent_ps,
                candidates_candidatesdetails.permanent_dist,
                employees_employeedetails.nominee_name,
                employees_employeedetails.nominee_father_name,
                employees_employeedetails.nominee_mohter_name,
                employees_employeedetails.nominee_vill,
                employees_employeedetails.nominee_po,
                employees_employeedetails.nominee_ps,
                employees_employeedetails.nominee_dist,
                employees_employeedetails.nominee_mobile_no,
                employees_employeedetails.relation_with_employee,
                employees_employeedetails.nominee_nid,
                employees_employeedetails.highest_degree,
  employees_employeedetails.profile_picture
FROM employees_employee
JOIN candidates_candidatesdetails ON employees_employee.candidate_id = candidates_candidatesdetails.candidate_id
JOIN employees_employeedetails ON employees_employee.id = employees_employeedetails.employee_id
JOIN employees_salaryinfo ON employees_employee.id = employees_salaryinfo.employee_id
join jobs_businessunit on employees_employee.unit_id = jobs_businessunit.id;
    '''
    with connection.cursor() as cursor:
        cursor.execute(query)
        results = cursor.fetchall()
        description = cursor.description
        if results != ():
            headers = [d[0] for d in description]
            output = Exporter().ExcelExporter(title='Employee List', headers=headers, body=results)
            output.seek(0)
            response = HttpResponse(output, content_type="application/vnd.ms-excel")
            response['content-Disposition'] = 'attachment; filename="Employee_Details.xlsx"'
            return response
        else:
            messages.warning(request, 'Empty List')


# Job Opening Report

def generate_jobs_pdf(request):
    if request.method == "POST":
        form = JobReportForm(request.POST)
        if form.is_valid():
            # Retrieve the units from the form
            unit = form.cleaned_data.get('unit')
            print(unit)
            
            # Convert unit to tuple to handle single and multiple units
            if len(unit) > 1:
                unit = tuple(unit)
            else:
                for i in unit:
                    unit = i
                unit = f'("{unit}")'
            

            # Prepare the query, dynamically inserting the units
            query = f'''
                SELECT j.posting_date AS Requisition_Date,
                    j.job_title AS POSITION,
                    j.department AS department,
                    j.job_location AS Location,
                    bu.short_name AS Unit,
                    DATEDIFF(CURDATE(), j.posting_date) AS "Lead Time",
                    j.no_of_position AS Total_Positions,
                    COUNT(o.id) AS Filled_Positions,
                    (j.no_of_position - COUNT(o.id)) AS "Vacant Positions"
                FROM jobs_job j
                JOIN jobs_businessunit bu ON j.unit_id = bu.id
                LEFT JOIN candidates_offer o ON o.job_id = j.id AND o.offer_status = 'Accepted'
                WHERE j.open_status = TRUE
                AND j.unit_id IN {unit}
                GROUP BY bu.short_name, j.posting_date, j.job_title, j.job_location, j.no_of_position, j.department
                ORDER BY bu.short_name, j.posting_date;
            '''

            # Execute the query and fetch the results
            with connection.cursor() as cursor:
                cursor.execute(query)
                result = cursor.fetchall()
                description = cursor.description
                headers = [d[0] for d in description]

                df = pd.DataFrame(result, columns=headers)

                # Create the pivot table
                df_pivot = pd.pivot_table(
                    df,
                    values='Vacant Positions',
                    index=['Unit', 'department', 'POSITION', 'Location', 'Lead Time'],
                    aggfunc='sum',
                    margins=True,
                    margins_name='Total',
                    fill_value=0
                )

                # Reset index and convert to list for PDF table
                sorted_df_pivot = df_pivot.sort_index()
                pivot_table = sorted_df_pivot.reset_index().values.tolist()

                # Prepare the table headers
                column_list = sorted_df_pivot.columns.to_list()
                column_list.insert(0, "Unit")
                column_list.insert(1, "Department")
                column_list.insert(2, "Position")
                column_list.insert(3, 'Location')
                column_list.insert(4, 'Lead Time (In Days)')

                pivot_table = [column_list] + pivot_table

                styles = getSampleStyleSheet()
                dept_summary_elements = []

                # Title style
                title_style = ParagraphStyle(
                    name='TitleStyle',
                    fontName="Times-Bold",
                    fontSize=18,
                    leading=20,
                    textColor=colors.black,
                    alignment=1,  # Center alignment
                )
                title = "<br/>Vacancy Report"
                dept_summary_elements.append(Paragraph(title, title_style))
                dept_summary_elements.append(Spacer(1, 15))

                # Convert text to paragraphs for word wrapping (for body rows)
                styleN = ParagraphStyle(
                    name='styleN',
                    fontSize=11,
                    leading=12,
                    fontName='Times-Roman',
                    textColor=colors.black
                )

                headerStyleN = ParagraphStyle(
                    name='styleN',
                    fontSize=11,
                    leading=12,
                    fontName='Times-Bold',
                    textColor=colors.black,
                    alignment=1
                )

                # Separate the header from the rest of the data
                header = pivot_table[0]  # First row (header)
                data = pivot_table[1:]  # The rest of the data

                # Wrap the data rows
                header_wrapped = [Paragraph(str(cell), headerStyleN) for cell in header]

                data_wrapped = [
                    [Paragraph(cell, styleN) if isinstance(cell, str) else cell for cell in row]
                    for row in data
                ]

                # Rebuild the table with the header and wrapped data
                final_table = [header_wrapped] + data_wrapped

                # Define the table and column widths
                dept_summary_table = Table(
                    final_table,
                    colWidths=[.8 * inch, 1.3 * inch, 1.8 * inch, 2.3 * inch, 0.9 * inch, 1 * inch]
                )

                # Get the span styles
                span_styles = auto_span_units_and_departments(pivot_table)

                # Define the table style, including bold for header
                dept_summary_table.setStyle(TableStyle([
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),  # Header background color
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),  # Bold font for the header
                    ('FONTNAME', (0, 1), (-1, -1), 'Times-Roman'),  # Normal font for body rows
                    ('FONTSIZE', (0, 0), (-1, 0), 12),  # Slightly larger font for header
                    ('FONTSIZE', (0, 1), (-1, -1), 11),  # Regular font size for body rows
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('TOPPADDING', (0, 0), (-1, 0), 6),
                    ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                    ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ALIGN', (2, 0), (3, -1), 'LEFT'),
                    ('SPAN', (0, -1), (4, -1)),
                ] + span_styles))

                # Add table to the elements
                
                dept_summary_elements.append(dept_summary_table)

                # Generate PDF with updated layout
                buffer = BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=20, leftMargin=20, topMargin=40, bottomMargin=40)
                doc.title = "Vacancy Report"
                doc.build(dept_summary_elements, onFirstPage=onFirstPage)

                # Send PDF as HTTP response
                buffer.seek(0)
                response = HttpResponse(buffer, content_type='application/pdf')
                response['Content-Disposition'] = 'inline; filename="Vacancy_Report.pdf"'
                return response
        else:
            messages.warning(request, "No units selected or form validation failed.")
            form = JobReportForm()
    else:
        form = JobReportForm()

    return render(request, "documents/jobVacancyReport.html", {'form': form})



def onFirstPage(canvas, doc):
    settings_dir = os.path.dirname(__file__)
    project_root = os.path.abspath(os.path.dirname(settings_dir))
    logo_path = os.path.join(project_root, 'logo.png')
    
    # Add logo to the first page
    canvas.drawImage(logo_path, 270, 780, width=35.62, height=45.32, mask='auto')


def auto_span_units_and_departments(pivot_table):
    unit_span_styles = []
    department_span_styles = []
    
    unit_col = [row[0] for row in pivot_table[1:]]  # Exclude header
    department_col = [row[1] for row in pivot_table[1:]]  # Exclude header

    start_unit_idx = 1
    start_dept_idx = 1
    
    # For unit spans
    for i in range(1, len(unit_col)):
        if unit_col[i] != unit_col[i - 1]:
            unit_span_styles.append(('SPAN', (0, start_unit_idx), (0, i)))  # Span Unit column
            start_unit_idx = i + 1
    unit_span_styles.append(('SPAN', (0, start_unit_idx), (0, len(unit_col))))  # Span last group for Unit

    # For department spans (nested under Unit)
    for i in range(1, len(department_col)):
        if department_col[i] != department_col[i - 1] or unit_col[i] != unit_col[i - 1]:
            department_span_styles.append(('SPAN', (1, start_dept_idx), (1, i)))  # Span Department column
            start_dept_idx = i + 1
    department_span_styles.append(('SPAN', (1, start_dept_idx), (1, len(department_col))))  # Span last group for Dept

    return unit_span_styles + department_span_styles


